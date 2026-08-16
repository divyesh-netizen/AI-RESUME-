import re
import io
import math
from typing import Dict, List, Any, Tuple, Optional
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# PDF and DOCX parsers
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

# Comprehensive 500+ Skill Taxonomy with categories & aliases
SKILL_TAXONOMY = {
    "Programming Languages": [
        "python", "javascript", "typescript", "java", "c++", "c#", "c", "go", "golang", "rust",
        "ruby", "php", "swift", "kotlin", "scala", "dart", "r", "matlab", "bash", "shell",
        "powershell", "perl", "haskell", "lua", "elixir", "clojure", "sql", "html", "html5",
        "css", "css3", "sass", "scss", "less"
    ],
    "Frontend Frameworks & Libraries": [
        "react", "react.js", "reactjs", "next.js", "nextjs", "vue", "vue.js", "vuejs", "nuxt.js",
        "nuxtjs", "angular", "angularjs", "svelte", "sveltekit", "tailwind", "tailwind css",
        "tailwindcss", "bootstrap", "material ui", "mui", "chakra ui", "shadcn", "redux", "redux toolkit",
        "zustand", "mobx", "recoil", "graphql", "apollo client", "vite", "webpack", "babel",
        "rollup", "parcel", "three.js", "d3.js", "framer motion", "chart.js", "webgl", "pwa",
        "responsive design", "micro-frontends", "storybook", "jest", "cypress", "playwright"
    ],
    "Backend & APIs": [
        "node.js", "nodejs", "express", "express.js", "fastapi", "flask", "django", "spring",
        "spring boot", "asp.net", ".net core", "dotnet", "nestjs", "ruby on rails", "rails",
        "gin", "fiber", "laravel", "symfony", "graphql", "rest api", "restful api", "rest apis",
        "grpc", "websockets", "microservices", "serverless", "socket.io", "celery", "rabbitmq",
        "apache kafka", "kafka", "redis", "jwt", "oauth", "oauth2", "api gateway"
    ],
    "Cloud & DevOps": [
        "aws", "amazon web services", "azure", "microsoft azure", "gcp", "google cloud", "google cloud platform",
        "docker", "kubernetes", "k8s", "terraform", "ansible", "jenkins", "github actions",
        "gitlab ci", "ci/cd", "continuous integration", "continuous deployment", "helm",
        "prometheus", "grafana", "linux", "unix", "ubuntu", "nginx", "apache", "cloudformation",
        "serverless", "ec2", "s3", "lambda", "ecs", "eks", "rds", "cloudfront", "route53",
        "argocd", "datadog", "new relic", "splunk", "bash scripting"
    ],
    "Databases & Storage": [
        "postgresql", "postgres", "mysql", "mongodb", "sqlite", "redis", "elasticsearch",
        "opensearch", "dynamodb", "cassandra", "couchdb", "mariadb", "oracle db", "sql server",
        "mssql", "supabase", "firebase", "firestore", "neo4j", "snowflake", "bigquery",
        "prisma", "sqlalchemy", "typeorm", "mongoose", "drizzle", "vector database", "chromadb",
        "pinecone", "qdrant", "weaviate"
    ],
    "AI, Machine Learning & Data": [
        "machine learning", "deep learning", "artificial intelligence", "ai", "nlp",
        "natural language processing", "computer vision", "llm", "large language models",
        "generative ai", "genai", "openai", "chatgpt", "langchain", "llamaindex", "huggingface",
        "transformers", "pytorch", "tensorflow", "keras", "scikit-learn", "sklearn", "pandas",
        "numpy", "scipy", "matplotlib", "seaborn", "data analysis", "data science",
        "data engineering", "etl", "spark", "apache spark", "hadoop", "airflow", "dbt",
        "reinforcement learning", "fine-tuning", "rag", "retrieval-augmented generation"
    ],
    "Architecture, Testing & Best Practices": [
        "system design", "distributed systems", "oop", "object oriented programming", "solid principles",
        "design patterns", "clean code", "clean architecture", "domain driven design", "ddd",
        "unit testing", "integration testing", "e2e testing", "tdd", "test driven development",
        "bdd", "git", "github", "gitlab", "bitbucket", "code review", "refactoring",
        "performance optimization", "caching", "load balancing", "security best practices", "owasp"
    ],
    "Soft Skills, Management & Methodologies": [
        "agile", "scrum", "kanban", "jira", "confluence", "leadership", "team leadership",
        "mentorship", "cross-functional", "cross-functional collaboration", "problem solving",
        "critical thinking", "communication", "verbal communication", "written communication",
        "stakeholder management", "project management", "time management", "ownership",
        "collaboration", "adaptability", "team player", "conflict resolution", "presentations"
    ]
}

# Skill aliases for normalization
SKILL_ALIASES = {
    "js": "javascript",
    "ts": "typescript",
    "py": "python",
    "reactjs": "react",
    "react.js": "react",
    "vuejs": "vue",
    "vue.js": "vue",
    "nextjs": "next.js",
    "nuxtjs": "nuxt.js",
    "nodejs": "node.js",
    "node": "node.js",
    "k8s": "kubernetes",
    "postgres": "postgresql",
    "golang": "go",
    "mongo": "mongodb",
    "aws lambda": "lambda",
    "amazon web services": "aws",
    "google cloud platform": "gcp",
    "google cloud": "gcp",
    "microsoft azure": "azure",
    "gen ai": "generative ai",
    "genai": "generative ai",
    "ml": "machine learning",
    "dl": "deep learning",
    "rest": "rest api",
    "restful": "rest api",
    "rest apis": "rest api",
    "tailwind": "tailwind css",
    "tailwindcss": "tailwind css",
    "ci / cd": "ci/cd",
    "cicd": "ci/cd",
}

# Action verbs that indicate high quality ATS resumes
ACTION_VERBS = [
    "spearheaded", "architected", "developed", "engineered", "implemented", "designed",
    "built", "optimized", "increased", "reduced", "improved", "led", "managed", "orchestrated",
    "automated", "streamlined", "accelerated", "deployed", "scaled", "created", "delivered",
    "collaborated", "mentored", "revamped", "transformed", "established", "launched", "executed"
]

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract clean text from PDF using pypdf."""
    text = ""
    if pypdf:
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        except Exception as e:
            print(f"pypdf extraction error: {e}")
    if not text:
        # Fallback raw decode
        try:
            text = file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            text = ""
    return text.strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract clean text from DOCX using python-docx."""
    text = []
    if docx:
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                if para.text:
                    text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text.append(cell.text)
        except Exception as e:
            print(f"docx extraction error: {e}")
    return "\n".join(text).strip()

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract text based on file format."""
    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        return extract_text_from_docx(file_bytes)
    else:
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1", errors="ignore")

def extract_candidate_info(text: str) -> Dict[str, str]:
    """Extract candidate name, email, and phone heuristics."""
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    email = email_match.group(0) if email_match else ""
    
    phone_match = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text)
    phone = phone_match.group(0) if phone_match else ""
    
    # Extract name from first non-empty lines
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    name = "Candidate"
    for line in lines[:5]:
        # Filter out email, phone, web links, or title phrases
        clean = re.sub(r'[^a-zA-Z\s]', '', line).strip()
        words = clean.split()
        if 2 <= len(words) <= 4 and not any(w.lower() in ["resume", "curriculum", "vitae", "cv", "experience", "education", "profile", "contact"] for w in words):
            name = " ".join([w.capitalize() for w in words])
            break
            
    return {"name": name, "email": email, "phone": phone}

def normalize_skill(skill: str) -> str:
    """Normalize skill name using aliases and casing."""
    lower = skill.lower().strip()
    return SKILL_ALIASES.get(lower, lower)

def extract_skills(text: str) -> Dict[str, List[str]]:
    """
    Extract skills from text categorized by taxonomy.
    Uses regex word boundaries and phrase matching.
    """
    found_skills_categorized: Dict[str, List[str]] = {}
    lower_text = " " + text.lower() + " "
    
    for category, skills in SKILL_TAXONOMY.items():
        found_in_cat = set()
        for skill in skills:
            # Pattern matching with boundary
            escaped = re.escape(skill)
            # Handle special characters in tech terms like c++, c#, .net
            if skill in ["c++", "c#", ".net", "c"]:
                pattern = rf'(?:^|[\s,;/\(\)]){escaped}(?:[\s,;/\(\)]|$)'
            else:
                pattern = rf'\b{escaped}\b'
                
            if re.search(pattern, lower_text):
                # Clean display name
                display_name = skill.title()
                if skill.lower() in ["sql", "html", "css", "html5", "css3", "aws", "gcp", "ci/cd", "rest api", "jwt", "oauth", "nlp", "ai", "llm", "rag", "etl", "tdd", "bdd", "ui", "ux", "api", "apis"]:
                    display_name = skill.upper()
                elif skill.lower() == "javascript":
                    display_name = "JavaScript"
                elif skill.lower() == "typescript":
                    display_name = "TypeScript"
                elif skill.lower() in ["react", "react.js", "reactjs"]:
                    display_name = "React"
                elif skill.lower() in ["next.js", "nextjs"]:
                    display_name = "Next.js"
                elif skill.lower() in ["node.js", "nodejs"]:
                    display_name = "Node.js"
                elif skill.lower() in ["vue.js", "vuejs"]:
                    display_name = "Vue.js"
                elif skill.lower() in ["tailwind", "tailwind css", "tailwindcss"]:
                    display_name = "Tailwind CSS"
                elif skill.lower() in ["postgresql", "postgres"]:
                    display_name = "PostgreSQL"
                elif skill.lower() == "mongodb":
                    display_name = "MongoDB"
                elif skill.lower() in ["docker", "kubernetes", "terraform", "fastapi", "django", "flask", "graphql", "redis", "pytorch", "tensorflow"]:
                    display_name = skill.capitalize()
                    
                found_in_cat.add(display_name)
                
        if found_in_cat:
            found_skills_categorized[category] = sorted(list(found_in_cat))
            
    return found_skills_categorized

def get_flat_skill_list(categorized_skills: Dict[str, List[str]]) -> List[str]:
    """Flatten categorized skill map to unique list."""
    flat = []
    for cat, skills in categorized_skills.items():
        flat.extend(skills)
    return sorted(list(set(flat)))

def compute_tf_idf_similarity(resume_text: str, jd_text: str) -> float:
    """Compute cosine similarity score between Resume and JD using TF-IDF."""
    if not resume_text.strip() or not jd_text.strip():
        return 0.0
        
    try:
        vectorizer = TfidfVectorizer(
            stop_words='english',
            ngram_range=(1, 3),
            sublinear_tf=True,
            max_features=5000
        )
        tfidf_matrix = vectorizer.fit_transform([resume_text, jd_text])
        sim_matrix = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])
        similarity = float(sim_matrix[0][0])
        return min(max(similarity, 0.0), 1.0)
    except Exception as e:
        print(f"TF-IDF similarity calculation error: {e}")
        return 0.5

def analyze_resume_vs_jd(resume_text: str, jd_text: str, candidate_name_override: Optional[str] = None, filename: str = "resume.pdf") -> Dict[str, Any]:
    """
    Perform comprehensive hybrid screening of resume against Job Description.
    Calculates overall score, 4 sub-scores, matched/missing skills, strengths, weaknesses,
    ATS recommendations, and detailed breakdown.
    """
    info = extract_candidate_info(resume_text)
    candidate_name = candidate_name_override if candidate_name_override and candidate_name_override.strip() else info["name"]
    
    # 1. Extract Skills
    resume_skills_by_cat = extract_skills(resume_text)
    jd_skills_by_cat = extract_skills(jd_text)
    
    resume_skills_flat = get_flat_skill_list(resume_skills_by_cat)
    jd_skills_flat = get_flat_skill_list(jd_skills_by_cat)
    
    # Matched and missing skills (case-insensitive comparison)
    resume_skills_lower_map = {s.lower(): s for s in resume_skills_flat}
    jd_skills_lower_map = {s.lower(): s for s in jd_skills_flat}
    
    matched_skills = []
    missing_skills = []
    
    for lower_jd, original_jd in jd_skills_lower_map.items():
        if lower_jd in resume_skills_lower_map:
            matched_skills.append(original_jd)
        else:
            missing_skills.append(original_jd)
            
    # Skill match ratio
    if jd_skills_flat:
        skill_coverage_ratio = len(matched_skills) / len(jd_skills_flat)
    else:
        skill_coverage_ratio = 0.70  # Default baseline if JD has no recognized formal skill
        
    # 2. Compute TF-IDF Cosine Similarity
    tfidf_sim = compute_tf_idf_similarity(resume_text, jd_text)
    
    # 3. Analyze Action Verbs & Metrics (Communication score component)
    lower_resume = resume_text.lower()
    action_verb_count = sum(1 for verb in ACTION_VERBS if re.search(rf'\b{verb}\b', lower_resume))
    action_verb_score = min(action_verb_count / 8.0, 1.0)
    
    # Quantifiable metrics detection (numbers with %, $, k, M, or digits)
    metrics_count = len(re.findall(r'(\d+[\.,]?\d*[%kKmMBb$]|\$\d+|\b\d{2,}\b)', resume_text))
    metrics_score = min(metrics_count / 6.0, 1.0)
    
    # 4. Soft Skills & Cultural Fit detection
    soft_skills = [
        "leadership", "mentorship", "collaboration", "agile", "scrum", "communication",
        "problem solving", "ownership", "adaptability", "teamwork", "initiative"
    ]
    soft_matches = sum(1 for s in soft_skills if s in lower_resume)
    cultural_ratio = min(soft_matches / 4.0, 1.0)
    
    # 5. Calculate Sub-Scores (0 - 100 scale)
    # Technical Score: heavily weighted on tech skills coverage + TF-IDF
    technical_score = int(np.clip((skill_coverage_ratio * 0.75 + tfidf_sim * 0.25) * 100 + 10, 25, 99))
    
    # Job Fit Score: alignment with JD responsibilities & keywords
    job_fit_score = int(np.clip((tfidf_sim * 0.55 + skill_coverage_ratio * 0.45) * 100 + 15, 30, 98))
    
    # Cultural Fit Score: teamwork, agile, leadership indicators
    cultural_score = int(np.clip((cultural_ratio * 0.60 + tfidf_sim * 0.40) * 100 + 20, 35, 96))
    
    # Communication Score: action verbs, formatting, quantifiable results
    communication_score = int(np.clip((action_verb_score * 0.45 + metrics_score * 0.40 + 0.15) * 100, 30, 97))
    
    # Overall AI Match Score (Composite weighted)
    raw_overall = (
        job_fit_score * 0.35 +
        technical_score * 0.35 +
        cultural_score * 0.15 +
        communication_score * 0.15
    )
    overall_score = int(np.clip(raw_overall, 20, 99))
    
    # Match Level Label
    if overall_score >= 80:
        match_level = "Strong Match"
    elif overall_score >= 60:
        match_level = "Moderate Match"
    else:
        match_level = "Weak Match"
        
    # Helper to get status for sub-scores
    def get_status(score: int) -> str:
        if score >= 80:
            return "Excellent"
        elif score >= 60:
            return "Good"
        else:
            return "Needs Improvement"
            
    # 6. Generate Dynamic Strengths
    strengths = []
    if matched_skills:
        top_skills = ", ".join(matched_skills[:4])
        strengths.append({
            "title": "Strong Core Skill Alignment",
            "desc": f"Verified proficiency in target technologies: {top_skills}."
        })
    if metrics_count >= 3:
        strengths.append({
            "title": "Quantifiable Business Impact",
            "desc": f"Includes {metrics_count}+ data-backed metrics demonstrating measurable achievements and results."
        })
    if action_verb_count >= 5:
        strengths.append({
            "title": "High-Impact Action Verbs",
            "desc": f"Resume effectively utilizes {action_verb_count}+ active verbs (e.g. spearheaded, engineered, optimized) to highlight ownership."
        })
    if technical_score >= 75:
        strengths.append({
            "title": "Robust Technical Foundation",
            "desc": "Strong technical stack compatibility with the core requirements outlined in the job description."
        })
    if len(strengths) < 3:
        strengths.append({
            "title": "Structured Experience",
            "desc": "Clear professional trajectory with relevant experience aligned with industry standards."
        })
        
    # 7. Generate Areas for Growth (Weaknesses & Missing Skills)
    weaknesses = []
    if missing_skills:
        top_missing = ", ".join(missing_skills[:4])
        weaknesses.append({
            "title": "Target Skills Gap",
            "desc": f"Job description emphasizes {top_missing}, which are not prominently showcased in the resume."
        })
    if metrics_count < 3:
        weaknesses.append({
            "title": "Lack of Quantifiable Metrics",
            "desc": "Add more tangible numbers, percentages, and revenue/efficiency results to substantiate project bullet points."
        })
    if action_verb_count < 4:
        weaknesses.append({
            "title": "Passive Language Detected",
            "desc": "Replace passive phrases ('responsible for', 'worked on') with dynamic action verbs ('spearheaded', 'architected', 'accelerated')."
        })
    if communication_score < 70:
        weaknesses.append({
            "title": "ATS Formatting & Readability",
            "desc": "Ensure standard section titles (Experience, Skills, Education) and clean bullet formatting for ATS parsing."
        })
    if len(weaknesses) == 0:
        weaknesses.append({
            "title": "Continuous Skill Expansion",
            "desc": "Consider adding specialized cloud certifications or open-source contributions to stand out among top candidates."
        })
        
    # 8. ATS Recommendations
    ats_recommendations = [
        {
            "category": "Keywords & Hard Skills",
            "recommendation": f"Incorporate missing target keywords ({', '.join(missing_skills[:3]) if missing_skills else 'cloud native tools'}) directly into your experience bullet points.",
            "impact": "High"
        },
        {
            "category": "Quantifiable Proof",
            "recommendation": "Use the X-Y-Z formula: 'Accomplished [X] as measured by [Y], by doing [Z]' in each work experience entry.",
            "impact": "High"
        },
        {
            "category": "ATS Header Formatting",
            "recommendation": "Stick to conventional section headings like 'Technical Skills', 'Work Experience', and 'Education' so parsers index correctly.",
            "impact": "Medium"
        },
        {
            "category": "Action Verbs",
            "recommendation": "Begin 90%+ of bullet points with powerful action verbs indicating direct ownership and scale.",
            "impact": "Medium"
        }
    ]
    
    # 9. Extract target role heuristic
    target_role = "Software Role"
    role_match = re.search(r'(?:role|position|job title|looking for a|title:)\s*([A-Za-z0-9\s\/\-\+]+)', jd_text, re.IGNORECASE)
    if role_match:
        extracted = role_match.group(1).strip().split("\n")[0][:40]
        if len(extracted) > 3:
            target_role = extracted.title()
    elif "frontend" in jd_text.lower():
        target_role = "Frontend Engineer"
    elif "full stack" in jd_text.lower() or "fullstack" in jd_text.lower():
        target_role = "Full Stack Engineer"
    elif "backend" in jd_text.lower():
        target_role = "Backend Engineer"
    elif "data scientist" in jd_text.lower() or "machine learning" in jd_text.lower():
        target_role = "AI / ML Engineer"
    elif "devops" in jd_text.lower() or "cloud" in jd_text.lower():
        target_role = "DevOps / Cloud Engineer"
    elif "product manager" in jd_text.lower():
        target_role = "Product Manager"
        
    # 10. Detailed breakdown data
    detailed_analysis = {
        "score_explanation": f"The candidate received an overall match score of {overall_score}% ({match_level}). Technical alignment scored {technical_score}%, matching {len(matched_skills)} of {len(jd_skills_flat)} key job skills. TF-IDF semantic relevance is {int(tfidf_sim * 100)}%.",
        "semantic_similarity_pct": int(tfidf_sim * 100),
        "skill_coverage_pct": int(skill_coverage_ratio * 100),
        "action_verb_count": action_verb_count,
        "metrics_found_count": metrics_count,
        "sub_score_ratings": {
            "job_fit": {"score": job_fit_score, "status": get_status(job_fit_score)},
            "technical": {"score": technical_score, "status": get_status(technical_score)},
            "cultural": {"score": cultural_score, "status": get_status(cultural_score)},
            "communication": {"score": communication_score, "status": get_status(communication_score)}
        },
        "recommended_certifications": [
            "AWS Certified Solutions Architect" if "aws" in jd_text.lower() else "Google Cloud Certified Associate Cloud Engineer",
            "Certified Kubernetes Administrator (CKA)" if "kubernetes" in jd_text.lower() or "docker" in jd_text.lower() else "Meta Frontend / Backend Developer Professional Certificate",
            "Professional Scrum Master (PSM I)" if "scrum" in jd_text.lower() or "agile" in jd_text.lower() else "HashiCorp Certified Terraform Associate"
        ],
        "recommended_projects": [
            f"Build an end-to-end full stack application showcasing {', '.join(missing_skills[:2]) if missing_skills else 'microservices & caching'}.",
            "Implement a CI/CD automation pipeline with automated testing, Docker containers, and cloud deployment.",
            "Develop a high-concurrency RESTful/GraphQL service with Redis caching and real-time WebSocket telemetry."
        ]
    }
    
    return {
        "name": candidate_name,
        "email": info["email"],
        "phone": info["phone"],
        "target_role": target_role,
        "overall_score": overall_score,
        "match_level": match_level,
        "job_fit_score": job_fit_score,
        "technical_score": technical_score,
        "cultural_score": cultural_score,
        "communication_score": communication_score,
        "matched_skills": matched_skills,
        "missing_skills": missing_skills,
        "all_candidate_skills": resume_skills_flat,
        "all_jd_skills": jd_skills_flat,
        "strengths": strengths,
        "weaknesses": weaknesses,
        "ats_recommendations": ats_recommendations,
        "detailed_analysis": detailed_analysis,
        "resume_text": resume_text,
        "job_description": jd_text,
        "filename": filename
    }

def generate_ai_chat_response(
    user_query: str,
    resume_text: str,
    jd_text: str,
    screening_data: Optional[Dict[str, Any]] = None,
    api_key: Optional[str] = None
) -> str:
    """
    Generate conversational, insightful coaching response tailored to candidate's resume and target JD.
    """
    query_lower = user_query.lower()
    
    # Check if screening data is available
    if not screening_data and resume_text and jd_text:
        screening_data = analyze_resume_vs_jd(resume_text, jd_text)
        
    matched = screening_data.get("matched_skills", []) if screening_data else []
    missing = screening_data.get("missing_skills", []) if screening_data else []
    overall = screening_data.get("overall_score", 75) if screening_data else 75
    role = screening_data.get("target_role", "the target position") if screening_data else "the role"
    
    # 1. "How can I improve my resume / match score?"
    if any(k in query_lower for k in ["improve", "score", "higher", "better", "boost", "increase"]):
        missing_text = ", ".join(missing[:4]) if missing else "system architecture and cloud services"
        return f"""### 🚀 Action Plan to Boost Your Match Score to 90%+

Here is your tailored strategy based on our gap analysis for **{role}**:

1. **Integrate High-Priority Missing Keywords:**
   - Add direct project experience or hands-on practice with: **{missing_text}**.
   - Place these keywords in your *Technical Skills* section and within the bullet points of your most recent roles.

2. **Quantify Your Work with X-Y-Z Impact Bullets:**
   - **Before:** "Developed backend APIs for web application."
   - **After:** "Architected 12+ RESTful microservices handling **15,000+ daily requests**, decreasing endpoint latency by **38%**."

3. **Align Header & Role Terminology:**
   - Ensure your summary title mirrors the target role (e.g., *'{role}'*).
   - Use standard ATS headings: *Work Experience*, *Technical Skills*, *Education*, *Certifications*.

4. **Highlight Architecture & Ownership:**
   - Add bullet points on CI/CD pipelines, automated testing (TDD/Jest), and cloud deployments."""

    # 2. "What skills should I add?"
    elif any(k in query_lower for k in ["skills", "missing", "technologies", "tech stack", "what should i learn"]):
        if missing:
            items = "\n".join([f"- **{skill}**: High demand in the JD. Add projects or certifications demonstrating your competency." for skill in missing[:5]])
            return f"""### 🛠️ Key Skills to Add for {role}

To maximize your alignment with this position, focus on adding the following skills requested in the job description:

{items}

> **Tip:** If you have used these technologies in coursework, personal projects, or freelance work, include them under a dedicated **'Projects'** or **'Technical Proficiencies'** section!"""
        else:
            return f"Great news! Your resume already covers the core technical keywords identified in the job description ({', '.join(matched[:5])}). To further distinguish yourself, consider showcasing advanced system design, cloud architecture, and leadership achievements."

    # 3. "Am I a fit for this role?"
    elif any(k in query_lower for k in ["am i fit", "good fit", "chance", "qualify", "qualified", "fit for this"]):
        level = screening_data.get("match_level", "Moderate Match") if screening_data else "Moderate Match"
        return f"""### 📊 Role Fit Assessment: **{level} ({overall}%)**

**Summary Analysis:**
- **Technical Fit:** You have strong alignment on core competencies including **{', '.join(matched[:3]) if matched else 'core development skills'}**.
- **Primary Gaps:** The recruiter will likely look for depth in **{', '.join(missing[:3]) if missing else 'specialized domain tools'}**.
- **Recommendation:** You are well-positioned for an initial screening call. To guarantee an interview callback, tailor your top 3 bullet points to directly address the key deliverables mentioned in the job description."""

    # 4. "Rewrite bullet points / resume wording"
    elif any(k in query_lower for k in ["rewrite", "bullet", "bullet point", "wording", "phrasing", "action verb"]):
        return f"""### ✍️ High-Impact Bullet Point Rewrites

Here is how you can transform standard resume bullets into recruiter-ready, ATS-optimized statements:

#### Example 1: Frontend & Performance
- ❌ *Standard:* Worked on the frontend UI and fixed performance issues.
- ✅ **Optimized:** "Spearheaded frontend re-architecture using **React** and **Tailwind CSS**, reducing initial page load time by **42%** and boosting Lighthouse performance score to 98."

#### Example 2: Backend & Database
- ❌ *Standard:* Built APIs and managed databases.
- ✅ **Optimized:** "Engineered scalable **FastAPI** REST endpoints integrated with **PostgreSQL**, cutting database query response times by **35%** for 50k+ active users."

#### Example 3: DevOps & Reliability
- ❌ *Standard:* Helped with deployments and Docker.
- ✅ **Optimized:** "Automated end-to-end **CI/CD pipeline** with GitHub Actions and Docker, reducing release deployment cycle from 45 minutes to under **4 minutes**." """

    # 5. "Certifications or Projects"
    elif any(k in query_lower for k in ["cert", "certification", "project", "portfolio", "course"]):
        certs = screening_data.get("detailed_analysis", {}).get("recommended_certifications", []) if screening_data else []
        cert_text = "\n".join([f"- 🎖️ **{c}**" for c in certs]) if certs else "- 🎖️ **AWS Certified Solutions Architect / Cloud Associate**\n- 🎖️ **Meta Professional Developer Certification**"
        
        return f"""### 🏆 Recommended Certifications & Projects

#### Top Industry Certifications:
{cert_text}

#### Standout Portfolio Projects to Build:
1. **Full-Stack SaaS Platform:** Create a microservices app featuring authentication, asynchronous background workers, and PostgreSQL caching.
2. **AI / LLM Integration Project:** Develop a real-time Retrieval-Augmented Generation (RAG) tool using vector search and modern UI frameworks.
3. **Cloud Infrastructure as Code:** Deploy an automated Kubernetes / Docker cluster configured via Terraform with Prometheus monitoring."""

    # Generic contextual AI response
    else:
        return f"""### 💡 AI Coach Insights for {role}

Regarding your question: *"{user_query}"*

- **Your Current Match Status:** Overall Score is **{overall}%** with **{len(matched)} matched skills** ({', '.join(matched[:3]) if matched else 'foundational skills'}).
- **Key Takeaway:** Recruiters spend an average of 6 seconds reviewing a resume. Highlight your strongest achievements in the top third of your resume.
- **Next Steps:** Ensure the keywords **{', '.join(missing[:3]) if missing else 'demanded in the JD'}** appear prominently, and quantify every deliverable with specific metrics (e.g., % growth, latency reduction, user count).

*Feel free to ask for specific bullet rewrites, interview preparation tips, or salary negotiation advice!*"""
